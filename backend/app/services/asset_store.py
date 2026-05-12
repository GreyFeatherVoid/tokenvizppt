import json
import mimetypes
import zipfile
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from fastapi import UploadFile
from openpyxl import load_workbook
from pypdf import PdfReader
from sqlalchemy import select

from app.core.settings import get_settings
from app.db.session import SessionLocal
from app.models.asset import Asset
from app.services.session_store import SessionNotFoundError, safe_session_id

ALLOWED_IMAGE_TYPES = {
    "image/jpeg": ".jpg",
    "image/png": ".png",
    "image/webp": ".webp",
    "image/gif": ".gif",
}
ALLOWED_DOCUMENT_TYPES = {
    "text/plain": ".txt",
    "text/markdown": ".md",
    "text/csv": ".csv",
    "application/csv": ".csv",
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
}
ALLOWED_IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".webp", ".gif"}
ALLOWED_DOCUMENT_SUFFIXES = {".txt", ".md", ".csv", ".pdf", ".docx", ".xlsx"}
MAX_IMAGE_SIZE = 8 * 1024 * 1024
MAX_DOCUMENT_SIZE = 20 * 1024 * 1024
MAX_EXTRACTED_TEXT_CHARS = 18000


class AssetValidationError(Exception):
    pass


class AssetNotFoundError(Exception):
    pass


class AssetStore:
    def __init__(self) -> None:
        self.storage_root = get_settings().storage_root
        self.assets_root = self.storage_root / "assets"
        self.assets_root.mkdir(parents=True, exist_ok=True)

    async def save_upload(self, session_id: str, upload: UploadFile) -> dict:
        session_id = safe_session_id(session_id)
        guessed_type = mimetypes.guess_type(upload.filename or "")[0] or ""
        content_type = upload.content_type or guessed_type
        if content_type not in {*ALLOWED_IMAGE_TYPES, *ALLOWED_DOCUMENT_TYPES} and guessed_type:
            content_type = guessed_type
        suffix_hint = Path(upload.filename or "").suffix.lower()
        suffix = self._resolve_suffix(content_type, upload.filename or "")

        data = await upload.read()
        if not data:
            raise AssetValidationError("Uploaded file is empty")
        asset_kind = (
            "image"
            if content_type in ALLOWED_IMAGE_TYPES or suffix_hint in ALLOWED_IMAGE_SUFFIXES
            else "document"
        )
        if asset_kind == "image" and len(data) > MAX_IMAGE_SIZE:
            raise AssetValidationError("Uploaded image exceeds 8MB")
        if asset_kind == "document" and len(data) > MAX_DOCUMENT_SIZE:
            raise AssetValidationError("Uploaded document exceeds 20MB")

        asset_id = uuid4().hex
        safe_name = Path(upload.filename or f"{asset_id}{suffix}").name
        relative_path = Path("assets") / session_id / f"{asset_id}{suffix}"
        absolute_path = self.storage_root / relative_path
        absolute_path.parent.mkdir(parents=True, exist_ok=True)
        absolute_path.write_bytes(data)

        metadata = {
            "kind": asset_kind,
            "notes": "",
            "required": False,
        }
        if asset_kind == "document":
            text = extract_document_text(data, content_type, safe_name)
            metadata.update(
                {
                    "text": text,
                    "text_char_count": len(text),
                }
            )

        row = Asset(
            id=asset_id,
            session_id=session_id,
            file_name=safe_name,
            file_path=relative_path.as_posix(),
            mime_type=content_type,
            file_size=len(data),
            metadata_json=json.dumps(metadata, ensure_ascii=False),
        )
        with SessionLocal() as db:
            db.add(row)
            db.commit()
            db.refresh(row)
            return self._to_dict(row)

    def save_generated_image(
        self,
        session_id: str,
        *,
        data: bytes,
        mime_type: str,
        metadata: dict,
    ) -> dict:
        session_id = safe_session_id(session_id)
        if not data:
            raise AssetValidationError("Generated image is empty")
        if mime_type not in ALLOWED_IMAGE_TYPES:
            raise AssetValidationError(f"Unsupported generated image type: {mime_type}")
        if len(data) > MAX_IMAGE_SIZE:
            raise AssetValidationError("Generated image exceeds 8MB")

        asset_id = uuid4().hex
        suffix = ALLOWED_IMAGE_TYPES[mime_type]
        file_name = f"ai-image-{asset_id[:12]}{suffix}"
        relative_path = Path("assets") / session_id / f"{asset_id}{suffix}"
        absolute_path = self.storage_root / relative_path
        absolute_path.parent.mkdir(parents=True, exist_ok=True)
        absolute_path.write_bytes(data)

        row_metadata = {
            "kind": "image",
            "notes": str(metadata.get("notes") or "").strip(),
            "required": False,
            "source": "ai_generated",
            "ai_image": metadata,
        }
        row = Asset(
            id=asset_id,
            session_id=session_id,
            file_name=file_name,
            file_path=relative_path.as_posix(),
            mime_type=mime_type,
            file_size=len(data),
            metadata_json=json.dumps(row_metadata, ensure_ascii=False),
        )
        with SessionLocal() as db:
            db.add(row)
            db.commit()
            db.refresh(row)
            return self._to_dict(row)

    def list_assets(self, session_id: str) -> list[dict]:
        session_id = safe_session_id(session_id)
        with SessionLocal() as db:
            rows = db.scalars(
                select(Asset)
                .where(Asset.session_id == session_id)
                .order_by(Asset.created_at.desc())
            ).all()
            return [self._to_dict(row) for row in rows]

    def get_asset(self, session_id: str, asset_id: str) -> dict:
        session_id = safe_session_id(session_id)
        with SessionLocal() as db:
            row = db.get(Asset, asset_id)
            if not row or row.session_id != session_id:
                raise AssetNotFoundError("Asset not found")
            return self._to_dict(row)

    def update_asset_metadata(
        self,
        session_id: str,
        asset_id: str,
        *,
        notes: str | None = None,
        required: bool | None = None,
    ) -> dict:
        session_id = safe_session_id(session_id)
        with SessionLocal() as db:
            row = db.get(Asset, asset_id)
            if not row or row.session_id != session_id:
                raise AssetNotFoundError("Asset not found")
            metadata = self._metadata(row)
            if notes is not None:
                metadata["notes"] = notes.strip()
            if required is not None:
                metadata["required"] = bool(required)
            row.metadata_json = json.dumps(metadata, ensure_ascii=False)
            db.commit()
            db.refresh(row)
            return self._to_dict(row)

    def update_asset_analysis(
        self,
        session_id: str,
        asset_id: str,
        analysis: dict,
    ) -> dict:
        session_id = safe_session_id(session_id)
        with SessionLocal() as db:
            row = db.get(Asset, asset_id)
            if not row or row.session_id != session_id:
                raise AssetNotFoundError("Asset not found")
            metadata = self._metadata(row)
            metadata["vision"] = analysis
            row.metadata_json = json.dumps(metadata, ensure_ascii=False)
            db.commit()
            db.refresh(row)
            return self._to_dict(row)

    def build_generation_context(self, session_id: str) -> dict:
        assets = self.list_assets(session_id)
        documents = []
        images = []
        for asset in assets:
            if asset["kind"] == "document" and asset.get("text"):
                documents.append(asset)
            elif asset["kind"] == "image":
                images.append(asset)
        return {
            "documents": documents,
            "images": images,
        }

    def get_asset_file_path(self, asset_id: str) -> Path:
        with SessionLocal() as db:
            row = db.get(Asset, asset_id)
            if not row:
                raise AssetNotFoundError("Asset not found")
            path = (self.storage_root / row.file_path).resolve()
            allowed_root = self.assets_root.resolve()
            if allowed_root not in [path, *path.parents] or not path.exists():
                raise SessionNotFoundError("Asset file not found")
            return path

    @staticmethod
    def _metadata(row: Asset) -> dict:
        try:
            metadata = json.loads(row.metadata_json or "{}")
        except json.JSONDecodeError:
            return {}
        return metadata if isinstance(metadata, dict) else {}

    @classmethod
    def _to_dict(cls, row: Asset) -> dict:
        metadata = cls._metadata(row)
        return {
            "id": row.id,
            "session_id": row.session_id,
            "file_name": row.file_name,
            "mime_type": row.mime_type,
            "file_size": row.file_size,
            "kind": metadata.get("kind") or "image",
            "source": metadata.get("source") or "uploaded",
            "notes": metadata.get("notes") or "",
            "required": bool(metadata.get("required")),
            "text": metadata.get("text") or "",
            "text_char_count": int(metadata.get("text_char_count") or 0),
            "vision": metadata.get("vision") if isinstance(metadata.get("vision"), dict) else {},
            "ai_image": metadata.get("ai_image")
            if isinstance(metadata.get("ai_image"), dict)
            else {},
            "url": f"/api/assets/{row.id}/file",
            "created_at": row.created_at.isoformat(),
        }

    @staticmethod
    def _resolve_suffix(content_type: str, filename: str) -> str:
        suffix = Path(filename).suffix.lower()
        if content_type in ALLOWED_IMAGE_TYPES:
            return ALLOWED_IMAGE_TYPES[content_type]
        if content_type in ALLOWED_DOCUMENT_TYPES:
            return ALLOWED_DOCUMENT_TYPES[content_type]
        if suffix in ALLOWED_IMAGE_SUFFIXES | ALLOWED_DOCUMENT_SUFFIXES:
            return suffix
        raise AssetValidationError(
            "Only jpg, png, webp, gif, txt, md, csv, pdf, docx, and xlsx files are supported"
        )


def extract_document_text(data: bytes, content_type: str, file_name: str) -> str:
    suffix = Path(file_name).suffix.lower()
    if content_type in {"text/plain", "text/markdown"} or suffix in {".txt", ".md"}:
        return _compact_text(data.decode("utf-8", errors="ignore"))
    if content_type in {"text/csv", "application/csv"} or suffix == ".csv":
        return _compact_text(data.decode("utf-8", errors="ignore"))
    if content_type == "application/pdf" or suffix == ".pdf":
        return _extract_pdf_text(data)
    if suffix == ".docx":
        return _extract_docx_text(data)
    if suffix == ".xlsx":
        return _extract_xlsx_text(data)
    return ""


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    chunks = [page.extract_text() or "" for page in reader.pages[:80]]
    return _compact_text("\n\n".join(chunks))


def _extract_docx_text(data: bytes) -> str:
    if not zipfile.is_zipfile(BytesIO(data)):
        raise AssetValidationError("Invalid docx file")
    document = Document(BytesIO(data))
    chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables[:20]:
        for row in table.rows[:80]:
            chunks.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return _compact_text("\n".join(chunks))


def _extract_xlsx_text(data: bytes) -> str:
    workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
    chunks = []
    for sheet in workbook.worksheets[:8]:
        chunks.append(f"Sheet: {sheet.title}")
        for row in sheet.iter_rows(max_row=120, max_col=20, values_only=True):
            values = [
                str(value).strip()
                for value in row
                if value is not None and str(value).strip()
            ]
            if values:
                chunks.append(" | ".join(values))
    workbook.close()
    return _compact_text("\n".join(chunks))


def _compact_text(value: str) -> str:
    text = "\n".join(line.strip() for line in value.replace("\r", "\n").split("\n"))
    text = "\n".join(line for line in text.split("\n") if line)
    return text[:MAX_EXTRACTED_TEXT_CHARS]


def get_asset_store() -> AssetStore:
    return AssetStore()
